"""
OCR / text extraction layer — zero hard dependencies, cascading fallbacks.

Fallback chains per file type:

  PDF (digital)  → pdfminer.six → pypdf → stdlib bytes scan
  PDF (scanned)  → PyMuPDF + (Tesseract → easyocr)  [caller may also use AI vision]
  Image          → Tesseract → easyocr               [caller may also use AI vision]
  DOCX           → python-docx → stdlib zipfile+xml  (zero deps, always works)
  DOC            → UTF-16 decode → olefile → antiword → LibreOffice → bytes scan
  TXT / CSV      → raw decode   (always works)
  Unknown        → probe by magic bytes, then try image/PDF paths

Nothing here blocks the module from loading — every optional library is imported
inside its own try/except so missing packages are skipped at runtime, not at startup.
"""
import base64
import io
import logging
import re
import subprocess
import xml.etree.ElementTree as ET
import zipfile

_logger = logging.getLogger(__name__)

# ── Install hints shown to users in error messages ────────────────────────────
_HINT_TESSERACT = (
    'For scanned PDFs and images, choose one option:\n'
    '\n'
    'Option A — pure Python (no system binary, any OS):\n'
    '  pip install easyocr\n'
    '\n'
    'Option B — Tesseract (higher speed, needs system binary):\n'
    '  Linux:   sudo apt install tesseract-ocr\n'
    '  macOS:   brew install tesseract\n'
    '  Windows: download installer from github.com/UB-Mannheim/tesseract/wiki\n'
    '  then:    pip install pytesseract Pillow PyMuPDF\n'
    '\n'
    'Option C — Cloud AI (zero extra packages, needs API key):\n'
    '  Set up Groq / OpenAI / Anthropic in Settings → Document Intelligence.'
)

# ── Tesseract → easyocr language code mapping ─────────────────────────────────
# Tesseract uses 3-letter codes (+joined), easyocr uses a list of 2-letter codes.
_TESS_TO_EASY = {
    'eng': 'en', 'fra': 'fr', 'deu': 'de', 'spa': 'es', 'por': 'pt',
    'ita': 'it', 'nld': 'nl', 'pol': 'pl', 'rus': 'ru', 'ara': 'ar',
    'chi_sim': 'ch_sim', 'chi_tra': 'ch_tra', 'jpn': 'ja', 'kor': 'ko',
    'kin': 'en',  # Kinyarwanda not in easyocr — fall back to English
    'swa': 'en',  # Swahili not in easyocr
}


def _tess_lang_to_easy(lang: str) -> list[str]:
    codes = [c.strip() for c in lang.split('+') if c.strip()]
    mapped = [_TESS_TO_EASY.get(c, 'en') for c in codes]
    return list(dict.fromkeys(mapped))  # deduplicate, preserve order


# ── Public entry point ─────────────────────────────────────────────────────────

def extract_text(file_data_b64: str, file_name: str, lang: str = 'eng') -> str:
    """
    Return extracted plain text from a base64-encoded file.

    :param file_data_b64: base64-encoded file bytes (Odoo Binary field value)
    :param file_name:     original filename — drives extractor selection
    :param lang:          Tesseract language code(s), e.g. 'eng' or 'eng+fra+kin'
    :raises RuntimeError: when no installed extractor can handle the file
    """
    raw_bytes = base64.b64decode(file_data_b64)
    fn = (file_name or '').lower()

    if fn.endswith('.pdf') or raw_bytes[:4] == b'%PDF':
        return _extract_pdf(raw_bytes, lang).strip()

    if fn.endswith('.docx'):
        return _extract_docx(raw_bytes).strip()

    if fn.endswith('.doc'):
        return _extract_doc(raw_bytes).strip()

    if fn.endswith(('.txt', '.csv', '.tsv', '.log')):
        return _decode_text(raw_bytes).strip()

    if fn.endswith(('.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.webp', '.gif')):
        text, err = _ocr_image_bytes(raw_bytes, lang)
        if text:
            return text.strip()
        raise RuntimeError(
            f'Cannot read this image without an OCR engine.\n{_HINT_TESSERACT}'
            + (f'\n\nDetail: {err}' if err else '')
        )

    # Unknown extension — probe by content
    return _probe_unknown(raw_bytes, lang, fn).strip()


# ── Public utility: render PDF pages to PNG base64 list ───────────────────────

def pdf_pages_to_png_b64(raw_bytes: bytes, dpi: int = 150, max_pages: int = 8) -> list[str]:
    """
    Render each PDF page to a PNG and return base64-encoded strings.

    Used by the AI vision OCR path in document_processor.py.
    Requires PyMuPDF (pip install PyMuPDF).
    Returns an empty list if PyMuPDF is not installed.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=raw_bytes, filetype='pdf')
        pages_b64 = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            pix = page.get_pixmap(dpi=dpi)
            png_bytes = pix.tobytes('png')
            pages_b64.append(base64.b64encode(png_bytes).decode('ascii'))
        _logger.info('Rendered %d PDF page(s) to PNG for vision OCR', len(pages_b64))
        return pages_b64
    except ImportError:
        _logger.debug('PyMuPDF not installed — cannot render PDF pages for vision OCR')
        return []
    except Exception as exc:
        _logger.warning('PDF page rendering failed: %s', exc)
        return []


# ── PDF ────────────────────────────────────────────────────────────────────────

def _extract_pdf(raw_bytes: bytes, lang: str) -> str:
    # ── Step 1: pure-Python text layer (handles most generated invoices) ──────
    text = _pdf_text_layer(raw_bytes)
    if text and len(text.strip()) > 40:
        return text

    # ── Step 2: page-by-page OCR for scanned PDFs ─────────────────────────────
    _logger.info('PDF text layer thin/empty — trying OCR per page')
    text, err = _pdf_ocr(raw_bytes, lang)
    if text and text.strip():
        return text

    # ── Step 3: give a clear, actionable error ─────────────────────────────────
    if err:
        raise RuntimeError(
            'This PDF appears to be scanned (image-only, no embedded text).\n'
            + _HINT_TESSERACT
            + f'\n\nDetail: {err}'
        )
    raise RuntimeError(
        'No text could be extracted from this PDF.\n' + _HINT_TESSERACT
    )


def _pdf_text_layer(raw_bytes: bytes) -> str:
    """Three pure-Python fallbacks — works on most digitally-generated PDFs."""

    # ── pdfminer.six (best quality, layout-aware) ─────────────────────────────
    try:
        from pdfminer.high_level import extract_text as _pdfminer
        from pdfminer.layout import LAParams
        # Tuned parameters: tighter word/line margins preserve table column separation
        params = LAParams(
            line_overlap=0.5,
            char_margin=2.0,
            line_margin=0.5,
            word_margin=0.1,
            boxes_flow=0.5,
            detect_vertical=False,
            all_texts=True,
        )
        text = _pdfminer(io.BytesIO(raw_bytes), laparams=params)
        if text and text.strip():
            _logger.info('pdfminer extracted %d chars', len(text))
            return text
    except ImportError:
        _logger.debug('pdfminer.six not installed')
    except Exception as exc:
        _logger.warning('pdfminer failed: %s', exc)

    # ── pypdf (pure Python, pip install pypdf, no binary) ─────────────────────
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw_bytes))
        pages = []
        for i, page in enumerate(reader.pages, 1):
            # extract_text with layout mode preserves column spacing better
            try:
                pt = page.extract_text(extraction_mode='layout') or ''
            except TypeError:
                pt = page.extract_text() or ''
            if pt.strip():
                pages.append(f'=== PAGE {i} ===\n{pt}')
        text = '\n\n'.join(pages)
        if text.strip():
            _logger.info('pypdf extracted %d chars across %d pages', len(text), len(pages))
            return text
    except ImportError:
        _logger.debug('pypdf not installed')
    except Exception as exc:
        _logger.warning('pypdf failed: %s', exc)

    # ── stdlib bytes scan (zero deps — works on many uncompressed PDFs) ────────
    try:
        raw_str = raw_bytes.decode('latin-1', errors='replace')
        chunks = re.findall(r'BT\s*(.*?)\s*ET', raw_str, re.DOTALL)
        words: list[str] = []
        for chunk in chunks:
            words += re.findall(r'\(([^)]{1,300})\)\s*Tj', chunk)
            words += re.findall(r'\(([^)]{1,300})\)', chunk)
        text = ' '.join(w for w in words if w.strip())
        if len(text.strip()) > 40:
            _logger.info('stdlib PDF bytes scan extracted %d chars', len(text))
            return text
    except Exception as exc:
        _logger.debug('stdlib PDF scan: %s', exc)

    return ''


def _pdf_ocr(raw_bytes: bytes, lang: str) -> tuple[str, str | None]:
    """Render PDF pages to images, then OCR. Returns (text, error_or_None)."""
    try:
        import fitz  # PyMuPDF — pure Python, no system binary: pip install PyMuPDF

        doc = fitz.open(stream=raw_bytes, filetype='pdf')
        pages_text: list[str] = []

        for i, page in enumerate(doc, 1):
            # 300 DPI gives Tesseract/easyocr enough resolution for small text
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes('png')

            page_text, _ = _ocr_image_bytes(img_bytes, lang)
            if page_text:
                pages_text.append(f'=== PAGE {i} ===\n{page_text}')

        if pages_text:
            text = '\n\n'.join(pages_text)
            _logger.info('PDF OCR extracted %d chars across %d page(s)', len(text), len(pages_text))
            return text, None

        return '', 'OCR produced no text — image may be too low resolution or blank'

    except ImportError as exc:
        return '', f'PyMuPDF not installed ({exc}). Run: pip install PyMuPDF'
    except Exception as exc:
        _logger.warning('PDF OCR failed: %s', exc)
        return '', str(exc)


# ── Image OCR (shared by image files and PDF page rendering) ──────────────────

def _preprocess_image_for_ocr(img_bytes: bytes) -> bytes:
    """
    Enhance image quality before OCR:
    - Upscale images smaller than 1000px wide to improve OCR on low-res scans
    - Convert to grayscale (better contrast for Tesseract/easyocr)
    - Apply contrast enhancement to make text stand out
    Returns processed bytes (PNG), or original bytes if PIL not available.
    """
    try:
        from PIL import Image, ImageEnhance, ImageFilter, UnidentifiedImageError
        try:
            img = Image.open(io.BytesIO(img_bytes))
        except (UnidentifiedImageError, Exception):
            return img_bytes

        # Upscale small images — OCR accuracy drops sharply below ~150 DPI equivalent
        w, h = img.size
        if w < 1000:
            scale = max(2, 1000 // w)
            img = img.resize((w * scale, h * scale), Image.LANCZOS)
            _logger.debug('Image upscaled %dx (was %dx%d)', scale, w, h)

        # Convert to grayscale for better OCR
        if img.mode not in ('L', 'RGB'):
            img = img.convert('RGB')
        gray = img.convert('L')

        # Moderate contrast boost — helps with faded or low-contrast scans
        enhancer = ImageEnhance.Contrast(gray)
        gray = enhancer.enhance(1.5)

        # Slight sharpening to clean up blurry scans
        gray = gray.filter(ImageFilter.SHARPEN)

        buf = io.BytesIO()
        gray.save(buf, format='PNG')
        return buf.getvalue()
    except ImportError:
        return img_bytes
    except Exception as exc:
        _logger.debug('Image preprocessing failed (non-fatal): %s', exc)
        return img_bytes


def _ocr_image_bytes(img_bytes: bytes, lang: str) -> tuple[str, str | None]:
    """
    OCR a raw image byte string. Tries Tesseract first, then easyocr.
    Returns (text, error_message_or_None).
    """
    # Preprocess image for better OCR accuracy
    processed = _preprocess_image_for_ocr(img_bytes)

    # ── Tesseract (fastest, most accurate when installed) ─────────────────────
    try:
        import pytesseract
        from PIL import Image, UnidentifiedImageError
        try:
            img = Image.open(io.BytesIO(processed))
        except UnidentifiedImageError:
            return '', 'Unrecognised image format'
        # OEM 3 = LSTM neural net, PSM 6 = assume uniform block of text
        config = '--oem 3 --psm 6'
        text = pytesseract.image_to_string(img, lang=lang, config=config)
        if text.strip():
            _logger.info('Tesseract OCR: %d chars', len(text))
            return text, None
        # Tesseract gave empty output — still try easyocr below
    except ImportError:
        _logger.debug('pytesseract not installed — trying easyocr')
    except Exception as exc:
        _logger.warning('Tesseract OCR failed: %s', exc)

    # ── easyocr (pure Python, no system binary: pip install easyocr) ──────────
    # First use downloads ML models (~100 MB). Subsequent calls are fast.
    try:
        import easyocr
        easy_langs = _tess_lang_to_easy(lang)
        _logger.info('easyocr attempt with langs=%s', easy_langs)
        reader = easyocr.Reader(easy_langs, verbose=False)
        # Pass preprocessed PNG bytes; easyocr accepts raw bytes directly
        results = reader.readtext(processed)
        text = '\n'.join(item[1] for item in results if item[1].strip())
        if text.strip():
            _logger.info('easyocr: %d chars', len(text))
            return text, None
    except ImportError:
        _logger.debug('easyocr not installed')
    except Exception as exc:
        _logger.warning('easyocr failed: %s', exc)

    return '', 'No OCR engine produced output'


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _extract_docx(raw_bytes: bytes) -> str:
    """Extract text from .docx. Falls back to stdlib if python-docx is missing."""

    # ── python-docx (richer: headers, footers, inline tables) ────────────────
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw_bytes))
        parts: list[str] = []

        # Headers and footers — company name and invoice number are often here
        for section in doc.sections:
            for hdr_para in section.header.paragraphs:
                t = hdr_para.text.strip()
                if t:
                    parts.append(t)
            for ftr_para in section.footer.paragraphs:
                t = ftr_para.text.strip()
                if t:
                    parts.append(t)

        # Body paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        # Tables — use tab separator so rule-based extractor can split columns
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                # Deduplicate merged cells (python-docx repeats merged cell text)
                deduped = []
                for cell in cells:
                    if not deduped or cell != deduped[-1]:
                        deduped.append(cell)
                if deduped:
                    parts.append('\t'.join(deduped))

        text = '\n'.join(parts)
        if text.strip():
            _logger.info('python-docx: %d chars', len(text))
            return text
    except ImportError:
        _logger.debug('python-docx not installed — using stdlib DOCX fallback')
    except Exception as exc:
        _logger.warning('python-docx failed: %s', exc)

    # ── stdlib fallback: DOCX is a ZIP containing word/document.xml ───────────
    # Works with ZERO packages — zipfile and xml.etree are Python stdlib.
    try:
        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
            with zf.open('word/document.xml') as xml_file:
                tree = ET.parse(xml_file)

        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        w = f'{{{ns}}}'
        parts = []

        for para in tree.iter(f'{w}p'):
            seg = ''.join(t.text or '' for t in para.iter(f'{w}t'))
            if seg.strip():
                parts.append(seg)

        text = '\n'.join(parts)
        if text.strip():
            _logger.info('stdlib DOCX fallback: %d chars', len(text))
            return text

    except KeyError:
        _logger.warning('word/document.xml not found in DOCX ZIP — file may be corrupted')
    except ET.ParseError as exc:
        _logger.warning('DOCX XML parse error: %s', exc)
    except Exception as exc:
        _logger.warning('stdlib DOCX fallback failed: %s', exc)

    raise RuntimeError(
        'Could not extract text from this DOCX file.\n'
        'For richer support: pip install python-docx'
    )


# ── DOC (legacy Word 97-2003 binary format) ────────────────────────────────────

def _extract_doc(raw_bytes: bytes) -> str:
    """
    Extract text from a legacy binary .doc file.
    Tries four approaches in order — at least the UTF-16 scan works with zero packages.
    """

    # ── 1. UTF-16 decode (zero deps — covers most Word 97-2003 documents) ─────
    # Word stores paragraph text as UTF-16LE strings inside the binary blob.
    try:
        raw_utf16 = raw_bytes.decode('utf-16-le', errors='ignore')
        # Strip NUL characters and unprintable control chars (keep \n \r \t)
        cleaned = re.sub(r'[^\x09\x0a\x0d\x20-\x7e\x80-\xff]', ' ', raw_utf16)
        cleaned = re.sub(r' {3,}', '  ', cleaned).strip()
        if len(cleaned) > 80:
            _logger.info('DOC UTF-16 scan: %d chars', len(cleaned))
            return cleaned
    except Exception as exc:
        _logger.debug('DOC UTF-16 scan failed: %s', exc)

    # ── 2. olefile (pure Python — pip install olefile) ────────────────────────
    # Reads OLE2 compound documents and extracts the WordDocument stream.
    try:
        import olefile
        with olefile.OleFileIO(io.BytesIO(raw_bytes)) as ole:
            if ole.exists('WordDocument'):
                stream = ole.openstream('WordDocument').read()
                # Word text is in UTF-16LE within the stream
                text = stream.decode('utf-16-le', errors='ignore')
                text = re.sub(r'[^\x09\x0a\x0d\x20-\x7e\x80-\xff]', ' ', text)
                text = re.sub(r' {3,}', '  ', text).strip()
                if len(text) > 80:
                    _logger.info('olefile DOC extraction: %d chars', len(text))
                    return text
    except ImportError:
        _logger.debug('olefile not installed')
    except Exception as exc:
        _logger.warning('olefile DOC extraction failed: %s', exc)

    # ── 3. antiword (system binary — apt install antiword) ────────────────────
    try:
        result = subprocess.run(
            ['antiword', '-'],
            input=raw_bytes, capture_output=True, timeout=15,
        )
        if result.returncode == 0 and result.stdout:
            text = result.stdout.decode('utf-8', errors='replace').strip()
            if text:
                _logger.info('antiword DOC extraction: %d chars', len(text))
                return text
    except FileNotFoundError:
        _logger.debug('antiword not found on PATH')
    except Exception as exc:
        _logger.warning('antiword failed: %s', exc)

    # ── 4. LibreOffice / soffice headless conversion .doc → text ─────────────
    # Command is 'libreoffice' on Linux/macOS, 'soffice' on some Windows installs.
    try:
        import os
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(raw_bytes)
            tmp_path = tmp.name
        try:
            out_dir = tempfile.mkdtemp()
            # Try both command names — whichever is on PATH wins
            lo_cmd = None
            for candidate in ('libreoffice', 'soffice'):
                try:
                    subprocess.run(
                        [candidate, '--version'],
                        capture_output=True, timeout=5,
                    )
                    lo_cmd = candidate
                    break
                except FileNotFoundError:
                    continue
            if lo_cmd:
                subprocess.run(
                    [lo_cmd, '--headless', '--convert-to', 'txt:Text',
                     '--outdir', out_dir, tmp_path],
                    capture_output=True, timeout=30,
                )
                # Build output path cross-platform (handles .DOC, .Doc, .doc)
                stem = os.path.splitext(os.path.basename(tmp_path))[0]
                txt_path = os.path.join(out_dir, stem + '.txt')
                if os.path.exists(txt_path):
                    with open(txt_path, encoding='utf-8', errors='replace') as f:
                        text = f.read().strip()
                    if text:
                        _logger.info('LibreOffice DOC conversion: %d chars', len(text))
                        return text
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except Exception as exc:
        _logger.warning('LibreOffice conversion failed: %s', exc)

    # ── 5. ASCII bytes scan (last resort) ─────────────────────────────────────
    try:
        strings = re.findall(r'[ -~\t]{20,}', raw_bytes.decode('latin-1', errors='replace'))
        text = '\n'.join(strings).strip()
        if len(text) > 80:
            _logger.info('DOC bytes scan (last resort): %d chars', len(text))
            return text
    except Exception as exc:
        _logger.debug('DOC bytes scan failed: %s', exc)

    raise RuntimeError(
        '.doc (legacy Word 97-2003 format) could not be extracted.\n'
        'Best fix: open in Word / LibreOffice, Save As .docx, and re-upload.\n'
        'Or install one of these (any OS):\n'
        '  pip install olefile       # pure Python, works on Windows/Mac/Linux\n'
        '  pip install python-docx   # sometimes handles .doc too\n'
        'System tools (if available):\n'
        '  antiword   — Linux/macOS: apt/brew install antiword\n'
        '  LibreOffice — all platforms: libreoffice.org'
    )


# ── Plain text / CSV ───────────────────────────────────────────────────────────

def _decode_text(raw_bytes: bytes) -> str:
    for enc in ('utf-8-sig', 'utf-8', 'utf-16', 'latin-1'):
        try:
            return raw_bytes.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return raw_bytes.decode('latin-1', errors='replace')


# ── Unknown file type — probe by magic bytes ───────────────────────────────────

def _probe_unknown(raw_bytes: bytes, lang: str, fn: str) -> str:
    # PDF magic
    if raw_bytes[:4] == b'%PDF':
        _logger.info('Detected PDF by magic bytes: %s', fn)
        return _extract_pdf(raw_bytes, lang)

    # ZIP-based (DOCX, XLSX, ODP, …)
    if raw_bytes[:2] == b'PK':
        _logger.info('Detected ZIP file (possible DOCX): %s', fn)
        try:
            return _extract_docx(raw_bytes)
        except RuntimeError:
            pass

    # OLE2 compound (legacy DOC, XLS, PPT)
    if raw_bytes[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
        _logger.info('Detected OLE2 compound file (possible .doc): %s', fn)
        try:
            return _extract_doc(raw_bytes)
        except RuntimeError:
            pass

    # Try image OCR
    text, _ = _ocr_image_bytes(raw_bytes, lang)
    if text and text.strip():
        return text

    # Try plain text
    try:
        return raw_bytes.decode('utf-8').strip()
    except UnicodeDecodeError:
        pass

    raise RuntimeError(
        f'Unsupported or unrecognised file format: "{fn}".\n'
        'Supported formats: PDF, DOCX, DOC, TXT, CSV, JPG, PNG, TIFF, BMP, WEBP.'
    )
